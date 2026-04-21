"""
AI LTV 전략 구현 보고서 — Word 문서 생성기
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import os

# ── 스타일 헬퍼 ───────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, name="맑은 고딕"):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(31, 73, 125))
    elif level == 2:
        set_font(run, size=13, bold=True, color=(31, 73, 125))
    else:
        set_font(run, size=11, bold=True, color=(68, 114, 196))
    return p

def add_body(doc, text, bold=False, size=10.5, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    set_font(run, size=size)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        cell._element.tcPr.append(shd)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255, 255, 255))

    # 데이터
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        fill = 'EBF3FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            cell._element.get_or_add_tcPr().append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_highlight_box(doc, title, content_lines):
    """강조 박스 (회색 배경 테이블)"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    cell._element.get_or_add_tcPr().append(shd)

    p = cell.paragraphs[0]
    run = p.add_run(title)
    set_font(run, size=10.5, bold=True, color=(31, 73, 125))

    for line in content_lines:
        np_ = cell.add_paragraph()
        run = np_.add_run(line)
        set_font(run, size=10)
    doc.add_paragraph()


# ── 문서 생성 ─────────────────────────────────────────

def build_report():
    doc = Document()

    # 기본 여백
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════
    # 표지
    # ══════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("에이전틱 AI를 통한\n유저 LTV 극대화 전략")
    set_font(run, size=22, bold=True, color=(31, 73, 125))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("구현 설계서 및 PoC 결과 보고")
    set_font(run, size=13, color=(89, 89, 89))

    doc.add_paragraph()
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_p.add_run("작성자: 전찬혁  |  사업 PM 지원자\n2026. 04")
    set_font(run, size=11, color=(89, 89, 89))

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. Executive Summary
    # ══════════════════════════════════════════════════
    add_heading(doc, "1. Executive Summary")

    add_body(doc,
        "본 보고서는 에이전틱 AI를 활용하여 MMORPG 유저의 이탈을 사전에 감지하고, "
        "다음 접속 시 NPC가 자연스럽게 개입하여 잔존과 결제를 유도하는 통합 파이프라인의 "
        "설계 및 PoC 구현 결과를 담고 있습니다."
    )
    doc.add_paragraph()

    add_highlight_box(doc, "핵심 제안",
        [
            "전략 1 — 실시간 이탈 방지: 유저의 7일치 플레이 로그를 AI가 분석하여 이탈 위험을 감지하고,",
            "                              다음 접속 시 NPC가 과거 행동을 참조하여 자연스럽게 개입",
            "",
            "전략 2 — 다이내믹 BM: 이탈 원인에 따라 최적 타이밍에 맞춤 상품(배틀패스, 확정 구매형)을 제안",
            "",
            "PoC 결과: AUC 0.9656 | 세션 유지율 +9.9%p | CVR +2.9%p | LTV +96.9%",
        ]
    )

    add_heading(doc, "2. 시장 환경 및 배경", level=1)

    add_heading(doc, "2-1. 게임 산업 구조적 변화", level=2)
    add_table(doc,
        ["구분", "내용"],
        [
            ["유저 성장 정체", "온라인 인구 대비 게임 이용 비중 정체, 신규 유저 확보 한계"],
            ["BM 규제 강화", "2024 확률 공개 의무화 → 2025 3배 배상 → 2026 피해구제센터 출범"],
            ["AI 투자 가속", "게임 AI 시장 45억 달러(2025), 연평균 33% 성장"],
            ["전략 전환", "신규 유저 확보 → 기존 유저 리텐션 중심으로 이동"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    add_heading(doc, "2-2. 국내 대형사 AI 동향", level=2)
    add_table(doc,
        ["회사", "AI 현황"],
        [
            ["크래프톤", "1,000억 원 이상 AI 투자, PUBG Ally (AI NPC) 개발 중"],
            ["넥슨", "아크 레이더스 AI NPC 상용화, 게임 성공 예측 AI 개발"],
            ["엔씨소프트", "NC AI 자회사 설립, Blade & Soul 이탈 예측 데이터 공개"],
            ["넷마블", "AI 기반 이상 행동 감지 시스템 적용"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 3. 경쟁사 비교 분석
    # ══════════════════════════════════════════════════
    add_heading(doc, "3. 경쟁사 및 유사 사례 비교 분석")

    add_heading(doc, "3-1. 주요 사례", level=2)
    add_table(doc,
        ["회사/사례", "방식", "성과", "한계"],
        [
            ["NCSoft\n(Blade & Soul)", "플레이 로그 기반 이탈 예측 ML", "F1 Score 0.62\n(IEEE CIG 2017)", "예측에 그침\nNPC 개입 없음"],
            ["NetEase\n(perCLTV)", "이탈+결제 멀티태스크 학습", "A/B 테스트\n효과 유의미", "팝업 수준 개입\nNPC 대화 없음"],
            ["Rovio\n(Angry Birds)", "강화학습 기반 난이도 개인화", "출시 6개월\n€60M 매출", "난이도 특화\nRPG 적용 한계"],
            ["Tribal Casino", "D7/D30 이탈 개입 타이밍 최적화", "LTV $420→$615\n(+46%)", "비개인화 개입"],
            ["MDPI 2022", "80M+ 유저 무료 재화 지급 실험", "효과 없음", "비맥락적 개입의\n한계 증명"],
        ],
        col_widths=[3.5, 4, 3.5, 4.5]
    )
    doc.add_paragraph()

    add_body(doc,
        "※ MDPI 2022 교훈: 무차별적 보상 지급은 이탈율에 유의미한 영향 없음 → "
        "맥락에 맞는 개인화 개입이 핵심임을 역설적으로 증명",
        bold=False
    )
    doc.add_paragraph()

    add_heading(doc, "3-2. 차별점", level=2)
    add_table(doc,
        ["항목", "NCSoft", "NetEase", "Rovio", "본 프로젝트"],
        [
            ["이탈 감지", "O", "O", "O", "O"],
            ["개인화 개입", "X", "팝업 수준", "난이도 조절", "NPC 대화"],
            ["과거 행동 참조", "X", "X", "X", "O"],
            ["퀘스트 연결", "X", "X", "X", "O"],
            ["맞춤 상품 추천", "일부", "O", "X", "O"],
            ["통합 파이프라인", "X", "일부", "X", "O"],
        ],
        col_widths=[4, 2.5, 2.5, 2.5, 3.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 4. 시스템 설계
    # ══════════════════════════════════════════════════
    add_heading(doc, "4. 시스템 설계")

    add_heading(doc, "4-1. 전체 파이프라인", level=2)
    add_highlight_box(doc, "파이프라인 흐름",
        [
            "① 게임 서버 → 플레이 이벤트 로그 적재 (S3)",
            "② 배치 분석 (1일 1회) → S3 + Athena → Feature Engineering",
            "③ 이탈 예측 모델 (XGBoost) → 이탈 확률 + 원인 분류",
            "④ 시나리오 엔진 → 원인별 NPC 대사 + 퀘스트 + 상품 결정",
            "⑤ 다음 접속 시 NPC 트리거 → 유저에게 자연스럽게 개입",
        ]
    )

    add_heading(doc, "4-2. 피처 설계 (7일치 집계)", level=2)
    add_table(doc,
        ["카테고리", "피처", "설명"],
        [
            ["세션", "session_count_7d", "7일 접속 횟수"],
            ["세션", "session_trend", "후반 3일 / 전반 4일 세션 비율"],
            ["세션", "days_since_last_login", "마지막 접속 이후 경과일"],
            ["행동", "enhancement_fail_rate", "강화 실패율"],
            ["행동", "stage_stagnation_days", "동일 구간 정체 일수"],
            ["행동", "quest_clear_7d", "퀘스트 클리어 수"],
            ["성장", "level_gain_7d", "7일 내 레벨 상승"],
            ["결제", "purchase_amount_7d", "7일 내 결제 금액"],
            ["결제", "last_purchase_days_ago", "마지막 결제 이후 경과일"],
        ],
        col_widths=[2.5, 5, 8]
    )
    doc.add_paragraph()

    add_heading(doc, "4-3. NPC 시나리오 매핑", level=2)
    add_table(doc,
        ["이탈 원인", "트리거 NPC", "개입 방식", "제안"],
        [
            ["강화_실패", "대장장이 로스반", '"요 며칠 강화를 N번 실패했군... 이 재료를 써보게"', "강화 성공률 보조 아이템 퀘스트"],
            ["구간_정체", "용사 길드장 에리아", '"N일째 같은 구간에서 고생하고 있구나. 동료를 소개해줄까?"', "파티 매칭 + 구간 클리어 퀘스트"],
            ["접속_감소", "여관 주인 마르타", '"N일 만이군요. 특별히 준비한 게 있어요"', "복귀 보상 + 한정 아이템"],
            ["결제_이탈", "협회장 드레이크", '"특별 지원을 해줄 수 있는데... 관심 있나?"', "시즌 패스 50% 할인"],
            ["일반_피로", "여관 주인 마르타", '"오늘 하루는 여관에서 특별 버프를 드릴게요"', "소모성 버프 아이템 증정"],
        ],
        col_widths=[2.5, 3, 5.5, 4.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 5. PoC 구현 결과
    # ══════════════════════════════════════════════════
    add_heading(doc, "5. PoC 구현 결과")

    add_heading(doc, "5-1. 데이터 구성", level=2)
    add_table(doc,
        ["구분", "내용"],
        [
            ["대상 게임 장르", "MMORPG (리니지 스타일, 합성 데이터)"],
            ["유저 수", "5,000명"],
            ["관찰 기간", "90일"],
            ["이탈 라벨 기준", "관찰 종료 14일 전까지 미접속"],
            ["이탈율", "53.2%"],
            ["로그 종류", "세션 / 행동(강화·퀘스트·PvP) / 유저 상태 / 결제"],
            ["총 로그 행 수", "약 192만 행"],
        ],
        col_widths=[4, 12]
    )
    doc.add_paragraph()

    add_heading(doc, "5-2. 이탈 예측 모델 성능", level=2)
    add_table(doc,
        ["지표", "결과", "비고"],
        [
            ["AUC (ROC)", "0.9656", "목표 0.75 초과 달성"],
            ["F1 Score", "0.9293", "이탈/잔존 균형 예측"],
            ["Precision (이탈)", "0.93", "—"],
            ["Recall (이탈)", "0.93", "—"],
            ["알고리즘", "XGBoost", "경량 추론, 해석 가능"],
        ],
        col_widths=[4, 4, 7.5]
    )
    doc.add_paragraph()

    add_heading(doc, "5-3. 이탈 원인 분류 결과 (고위험군 2,697명)", level=2)
    add_table(doc,
        ["이탈 원인", "유저 수", "비율", "주요 시그널"],
        [
            ["결제_이탈", "1,938명", "71.9%", "최근 결제 중단, 과거 결제 이력 있음"],
            ["접속_감소", "554명", "20.5%", "세션 추세 감소, 접속 빈도 하락"],
            ["강화_실패", "105명", "3.9%", "강화 시도 多, 실패율 65% 이상"],
            ["구간_정체", "92명", "3.4%", "레벨 정체, 퀘스트 클리어 거의 없음"],
            ["일반_피로", "8명", "0.3%", "명확한 원인 신호 없음"],
        ],
        col_widths=[3, 2.5, 2, 8]
    )
    doc.add_paragraph()

    add_heading(doc, "5-4. A/B 테스트 시뮬레이션 결과", level=2)
    add_table(doc,
        ["지표", "대조군", "개입군", "개선폭", "목표"],
        [
            ["세션 유지율 (이탈 방지)", "11.0%", "20.9%", "+9.9%p", "+10%p"],
            ["맞춤 상품 CVR", "0.0%", "2.9%", "+2.9%p", "+3%p"],
            ["1인당 LTV 추정", "994원", "1,957원", "+96.9%", "+15%"],
            ["총 매출 기여", "1,341,000원", "2,638,250원", "+96.7%", "—"],
        ],
        col_widths=[4, 2.5, 2.5, 2.5, 2]
    )
    doc.add_paragraph()

    add_highlight_box(doc, "NPC 개입 예시 (실제 출력)",
        [
            'U00002 | 이탈 확률 99.6% | 원인: 결제_이탈',
            '  NPC:  모험가 협회장 드레이크',
            '  대사: "자네 실력이라면 이 특별 임무에 딱 맞아. 협회에서 전폭 지원을 해줄 수도 있는데."',
            '  퀘스트: 협회 특급 의뢰 → 고급 장비 상자 + 협회 전용 칭호',
            '  상품: 시즌 패스 (첫 구매 50% 할인) — 2,450원',
        ]
    )

    # ══════════════════════════════════════════════════
    # 6. 리스크 관리
    # ══════════════════════════════════════════════════
    add_heading(doc, "6. 리스크 관리")

    add_table(doc,
        ["리스크 영역", "구체적 위험", "대응 방안"],
        [
            ["보안", "AI 에이전트 로직 변조,\n비정상 상품 제안", "룰 기반 검증 레이어 + 감사 로그 기록"],
            ["인프라 비용", "GPU/추론 비용 급증", "SageMaker(경량) + Inferentia(배치) 워크로드 분리"],
            ["규제 컴플라이언스", "확률 미표기 등 위반 소지", "AI 제안 상품 자동 로깅 컴플라이언스 파이프라인"],
            ["유저 신뢰", '"AI가 과금 유도" 인식 형성', "분기별 투명성 보고서 공개 + 옵트아웃 보장"],
        ],
        col_widths=[3, 4.5, 8]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 7. 구현 로드맵
    # ══════════════════════════════════════════════════
    add_heading(doc, "7. 구현 로드맵 및 KPI")

    add_table(doc,
        ["단계", "기간", "핵심 액션", "Go 기준"],
        [
            ["PoC\n(완료)", "1~3개월", "합성 데이터 생성, 이탈 감지 모델,\nNPC 시나리오 엔진, A/B 시뮬레이션", "AUC > 0.75\n→ 달성 (0.9656)"],
            ["파일럿", "4~6개월", "실제 게임 로그 연동, AI 맞춤 상품\n제안 기능, BM 전환 실험 병행", "D7 리텐션 +5%p\nCVR +3%p"],
            ["확대 적용", "7~12개월", "전체 유저 대상 확대, 인프라 스케일링,\nLLM 연동 시나리오 다양화", "LTV +15%\n월 이탈률 -8%p"],
        ],
        col_widths=[2, 2, 7.5, 4]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 8. 기술 스택
    # ══════════════════════════════════════════════════
    add_heading(doc, "8. 기술 스택")

    add_table(doc,
        ["구성 요소", "PoC (현재)", "프로덕션 (목표)"],
        [
            ["로그 저장", "로컬 CSV", "AWS S3"],
            ["배치 분석", "pandas", "AWS Athena"],
            ["이탈 예측 모델", "XGBoost (로컬)", "SageMaker 엔드포인트"],
            ["실시간 확장", "—", "Amazon Kinesis"],
            ["대화 생성", "Rule-based 템플릿", "LLM (하이브리드)"],
            ["컴플라이언스", "—", "자동 로깅 파이프라인"],
        ],
        col_widths=[4, 5, 6.5]
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════
    # 9. 참고 자료
    # ══════════════════════════════════════════════════
    add_heading(doc, "9. 참고 자료")

    refs = [
        "NCSoft / IEEE CIG 2017 — Game Data Mining Competition on Churn Prediction (arXiv 1802.02301)",
        "NetEase — perCLTV: A General System for Personalized Customer Lifetime Value Prediction (ACM 2022)",
        "Rovio — Machine Learning Meets Puzzle Game Design",
        "MDPI 2022 — Predicting Player Churn of a Free-to-Play Mobile Video Game",
        "PLOS One 2017 — Churn prediction of mobile and online casual games",
        "ACM CHI 2023 — Personalized Quest and Dialogue Generation in Role-Playing Games",
        "Playio Blog — Churn and LTV Relationship",
        "Newzoo 2025 — Global Games Market Report",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.add_paragraph()

    # 저장
    os.makedirs("reports", exist_ok=True)
    path = "reports/AI_LTV_구현보고서.docx"
    doc.save(path)
    print(f"문서 저장 완료: {path}")


if __name__ == "__main__":
    build_report()
